from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_dummy_cv_docx(filename="dummy_cv.docx"):
    document = Document()

    # Title
    document.add_heading("John Doe", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Information
    document.add_heading("Contact Information", level=2)
    document.add_paragraph("Email: john.doe@example.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe")
    document.add_paragraph("GitHub: github.com/johndoe | Portfolio: johndoe.dev")
    document.add_paragraph() # Spacer

    # Summary (Optional, but good to have)
    document.add_heading("Summary", level=2)
    document.add_paragraph(
        "A highly motivated and skilled software engineer with experience in Python, Java, and web development. "
        "Proven ability to lead projects and contribute to team success."
    )
    document.add_paragraph()

    # Skills
    document.add_heading("Skills", level=2)
    document.add_paragraph("Programming Languages: Python, Java, C++, JavaScript, SQL")
    document.add_paragraph("Frameworks & Technologies: Spring Boot, Django, React, Node.js")
    document.add_paragraph("Tools: Git, Docker, Kubernetes, Jenkins, AWS")
    document.add_paragraph("Databases: PostgreSQL, MongoDB, Redis")
    document.add_paragraph()

    # Work Experience
    document.add_heading("Work Experience", level=2)
    # Job 1
    document.add_paragraph("Senior Software Engineer", style='Intense Quote') # Using a style to make it stand out
    p = document.add_paragraph()
    p.add_run("Tech Solutions Inc.").bold = True
    p.add_run(" | Jan 2021 - Present")
    document.add_paragraph("- Developed new features for the company's flagship product using Python and Django.")
    document.add_paragraph("- Led a scrum team of 5 engineers, overseeing project planning and execution.")
    document.add_paragraph("- Mentored junior developers and conducted code reviews.")
    document.add_paragraph()
    
    # Job 2
    document.add_paragraph("Software Engineer", style='Intense Quote')
    p = document.add_paragraph()
    p.add_run("Innovate Corp.").bold = True
    p.add_run(" | Jun 2019 - Dec 2020")
    document.add_paragraph("- Contributed to the development of a microservices-based application using Java and Spring Boot.")
    document.add_paragraph("- Assisted with testing, bug fixing, and improving application performance.")
    document.add_paragraph()

    # Education
    document.add_heading("Education", level=2)
    # Degree 1
    document.add_paragraph("Master of Science in Computer Science", style='Intense Quote')
    p = document.add_paragraph()
    p.add_run("XYZ University").bold = True
    p.add_run(" | Graduated: May 2020")
    document.add_paragraph("Relevant Coursework: Data Structures, Advanced Algorithms, Machine Learning.")
    document.add_paragraph("Thesis: Applied Machine Learning to XYZ problem.")
    document.add_paragraph()

    # Degree 2
    document.add_paragraph("Bachelor of Science in Software Engineering", style='Intense Quote')
    p = document.add_paragraph()
    p.add_run("ABC College").bold = True
    p.add_run(" | Graduated: Dec 2018")
    document.add_paragraph("Minor: Mathematics")
    document.add_paragraph("Capstone Project: Developed a mobile application for campus navigation.")
    document.add_paragraph()

    # Projects
    document.add_heading("Projects", level=2)
    # Project 1
    document.add_paragraph("Personal Portfolio Website", style='Intense Quote')
    p = document.add_paragraph()
    p.add_run("Tech Stack:").bold = True
    p.add_run(" HTML, CSS, JavaScript, Flask (Python)")
    document.add_paragraph(
        "- Developed and deployed a responsive personal portfolio website to showcase projects and skills."
    )
    document.add_paragraph("Link: github.com/johndoe/portfolio")
    document.add_paragraph()

    # Project 2
    document.add_paragraph("AI Powered Recommendation System", style='Intense Quote')
    p = document.add_paragraph()
    p.add_run("Tech Stack:").bold = True
    p.add_run(" Python, Pandas, Scikit-learn")
    document.add_paragraph("- Built a content-based recommendation system for movies using collaborative filtering techniques.")
    
    document.save(filename)
    print(f"'{filename}' created successfully.")

if __name__ == "__main__":
    create_dummy_cv_docx()
