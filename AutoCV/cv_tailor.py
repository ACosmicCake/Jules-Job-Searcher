# cv_tailor.py
# Main script for the AI-powered CV tailoring program.

import os
from dotenv import load_dotenv
import sys
import json
from google import genai
from PyPDF2 import PdfReader
import docx
from generate_cv import generate_cv_pdf_from_json_string

def get_api_key() -> str | None:
    """Loads the Google API key from environment variables or .env file."""
    load_dotenv()
    api_key = os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        print("Error: GOOGLE_API_KEY not found.")
        print("Please set it as an environment variable or in a .env file.")
        sys.exit(1) # Exit if key is not found
    return api_key

def get_cv_from_text_file(filepath: str) -> str | None:
    """Reads CV content from a plain text file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: CV file not found at {filepath}")
        return None

def get_cv_from_json_file(filepath: str) -> dict | None:
    """Reads CV content from a JSON file and parses it into a dictionary."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f) # Parses JSON into a Python dictionary
    except FileNotFoundError:
        print(f"Error: CV JSON file not found at {filepath}")
        return None
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON from {filepath}. Please ensure it's valid JSON.")
        return None

def get_cv_from_user_paste() -> str:
    """Gets CV content by asking the user to paste it into the console."""
    print("\nPaste your full CV content below. Press Ctrl+D (Unix/Linux) or Ctrl+Z then Enter (Windows) when done:")
    lines = []
    while True:
        try:
            line = input()
        except EOFError: # Detects Ctrl+D or Ctrl+Z
            break
        lines.append(line)
    return "\n".join(lines)

def get_job_description_from_user_paste() -> str:
    """Gets job description by asking the user to paste it into the console."""
    print("\nPaste the job description below. Press Ctrl+D (Unix/Linux) or Ctrl+Z then Enter (Windows) when done:")
    lines = []
    while True:
        try:
            line = input()
        except EOFError: # Detects Ctrl+D or Ctrl+Z
            break
        lines.append(line)
    return "\n".join(lines)

def get_job_description_from_file(filepath: str) -> str | None:
    """Reads job description content from a plain text file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: Job description file not found at {filepath}")
        return None

def call_gemini_api(api_key: str, prompt_text: str) -> str | None:
    """
    Calls the Gemini API with the provided prompt and API key.

    Args:
        api_key: The Google API key.
        prompt_text: The complete prompt to send to the model.

    Returns:
        The AI-generated text response, or None if an error occurs.
    """
    try:
        client = genai.Client(api_key=api_key)
        model_to_use = "gemini-2.5-flash-preview-05-20" # Or other compatible model like "gemini-1.0-pro"

        response = client.models.generate_content(
            model = model_to_use,
            contents=prompt_text
        )
        # Accessing the text response, ensuring parts and text exist
        if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
            return response.candidates[0].content.parts[0].text
        else:
            # Handle cases where the response structure is unexpected or content is missing
            print("Warning: Gemini API response structure was not as expected or content was empty.")
            # Attempt to access response.text directly as a fallback
            if hasattr(response, 'text'):
                return response.text
            return None
    except Exception as e:
        print(f"Error calling Gemini API (google-genai): {e}")
        return None

def get_cv_from_pdf_file(filepath: str) -> str | None:
    """
    Extracts text content from a PDF file.

    Args:
        filepath: The path to the PDF file.

    Returns:
        A string containing the extracted text, or None if an error occurs.
    """
    try:
        text_content = []
        with open(filepath, 'rb') as f: # Open in binary read mode
            reader = PdfReader(f)
            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception as decrypt_error:
                    print(f"Error: Could not decrypt PDF {filepath}. It might be password-protected. Error: {decrypt_error}")
                    return None

            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text: # Ensure text was extracted
                    text_content.append(page_text)

        if not text_content:
            print(f"Warning: No text could be extracted from PDF {filepath}. It might be an image-based PDF or have non-standard text encoding.")
            return ""

        return "\n".join(text_content)
    except FileNotFoundError:
        print(f"Error: PDF file not found at {filepath}")
        return None
    except Exception as e:
        print(f"Error processing PDF file {filepath}: {e}")
        return None

def get_cv_from_docx_file(filepath: str) -> str | None:
    """
    Extracts text content from a DOCX file.

    Args:
        filepath: The path to the DOCX file.

    Returns:
        A string containing the extracted text, or None if an error occurs.
    """
    try:
        document = docx.Document(filepath)
        text_content = []
        for para in document.paragraphs:
            text_content.append(para.text)

        if not text_content and not document.tables:
            print(f"Warning: No text paragraphs could be extracted from DOCX {filepath}. It might be empty or contain non-standard content.")
            return ""

        return "\n".join(text_content)
    except FileNotFoundError:
        print(f"Error: DOCX file not found at {filepath}")
        return None
    except Exception as e:
        print(f"Error processing DOCX file {filepath}: {e}")
        return None

def main():
    print("--- AI-Powered CV Tailoring Program ---")
    api_key = get_api_key()
    if not api_key:
        return
    print("Successfully loaded API key.")

    cv_data = None # Stores raw CV data (dict if json, str otherwise)
    job_description_text = None
    cv_data_for_prompt = None # Stores CV data as a string for the prompt

    # 1. Get CV Data
    print("\n--- Step 1: Provide Your CV ---")
    cv_input_choice = input("How would you like to provide your CV? (json / text / pdf / docx / paste / skip): ").strip().lower()
    if cv_input_choice == 'json':
        cv_filepath = input("Enter path to your CV JSON file (e.g., my_cv.json): ").strip()
        cv_data = get_cv_from_json_file(cv_filepath)
        if cv_data:
            print(f"Successfully loaded CV from {cv_filepath}")
            cv_data_for_prompt = json.dumps(cv_data, indent=2) # Convert dict to formatted JSON string
    elif cv_input_choice == 'text':
        cv_filepath = input("Enter path to your CV text file (e.g., my_cv.txt): ").strip()
        cv_data_for_prompt = get_cv_from_text_file(cv_filepath) # Directly use string content
        if cv_data_for_prompt is not None:
            print(f"Successfully loaded CV from {cv_filepath}")
    elif cv_input_choice == 'pdf':
        cv_filepath = input("Enter path to your CV PDF file (e.g., my_cv.pdf): ").strip()
        cv_data_for_prompt = get_cv_from_pdf_file(cv_filepath)
        if cv_data_for_prompt is not None:
            print(f"Successfully extracted text from PDF CV at {cv_filepath}")
            if not cv_data_for_prompt:
                 print("Warning: The extracted text from the PDF is empty.")
    elif cv_input_choice == 'docx':
        cv_filepath = input("Enter path to your CV DOCX file (e.g., my_cv.docx): ").strip()
        cv_data_for_prompt = get_cv_from_docx_file(cv_filepath)
        if cv_data_for_prompt is not None:
            print(f"Successfully extracted text from DOCX CV at {cv_filepath}")
            if not cv_data_for_prompt:
                 print("Warning: The extracted text from the DOCX is empty.")
    elif cv_input_choice == 'paste':
        cv_data_for_prompt = get_cv_from_user_paste()
        if cv_data_for_prompt:
            print("Successfully received CV via paste.")
        else:
            print("No CV content was pasted.")
    elif cv_input_choice == 'skip':
        print("CV input skipped.")
    else:
        print("Invalid choice for CV input. Skipping.")

    # 2. Get Job Description
    print("\n--- Step 2: Provide Job Description ---")
    jd_input_choice = input("How would you like to provide the Job Description? (paste / text / skip): ").strip().lower()
    if jd_input_choice == 'paste':
        job_description_text = get_job_description_from_user_paste()
        if job_description_text:
            print("Successfully received job description via paste.")
    elif jd_input_choice == 'text':
        jd_filepath = input("Enter path to the job description text file (e.g., job_description.txt): ").strip()
        job_description_text = get_job_description_from_file(jd_filepath)
        if job_description_text:
            print(f"Successfully loaded job description from {jd_filepath}")
    elif jd_input_choice == 'skip':
        print("Job description input skipped.")
    else:
        print("Invalid choice for job description input. Skipping.")

    if cv_data_for_prompt and job_description_text:
        print("\n--- Step 3: Processing ---")

        # Load CV template structure from CV_format.json
        cv_template_content = ""  # Default to empty string
        cv_format_filename = "CV_format.json"
        try:
            with open(cv_format_filename, 'r', encoding='utf-8') as f:
                cv_template_content = f.read()
            if not cv_template_content.strip():
                print(f"Warning: '{cv_format_filename}' was found but is empty or contains only whitespace. Using an empty structure for CV_template.")
                cv_template_content = ""  # Ensure it's truly empty if only whitespace
            else:
                print(f"Successfully loaded CV labeling structure from '{cv_format_filename}'.")
        except FileNotFoundError:
            print(f"Warning: '{cv_format_filename}' not found. The CV labeling structure part of the prompt will be empty.")
        except Exception as e:
            print(f"Error reading '{cv_format_filename}': {e}. The CV labeling structure part of the prompt will be empty.")

        prompt_text = f"""
You are an expert CV tailoring assistant. Your task is to rewrite the provided CV to be perfectly tailored for the given job description.

Follow these instructions carefully:
1.  Analyze the job description for key skills, experience, and keywords.
2.  Rewrite the CV's summary and work experience sections to highlight these aspects.
3.  Use strong action verbs and quantify achievements where possible.
4.  Ensure the tone is professional and matches the industry.
5.  The output should be a complete, well-formatted CV. Do not output anything else before or after the CV content itself.

Here is the original CV:
--- BEGIN CV ---
{cv_data_for_prompt}
--- END CV ---

Here is the target job description:
--- BEGIN JOB DESCRIPTION ---
{job_description_text}
--- END JOB DESCRIPTION ---

Generate the text based on this labeling structure:
--- BEGIN LABELING STRUCTURE ---
{cv_template_content}
--- END LABELING STRUCTURE ---

Now, please provide the tailored CV. If there is nothing in the structure place , do not output anthing in that spot:
"""
        print(f"Generated prompt (first 200 chars): {prompt_text[:200].replace(os.linesep, ' ')}...") # Show more context

        print("Calling Gemini API...")
        tailored_cv_output = call_gemini_api(api_key, prompt_text)

        if tailored_cv_output:
            # Clean the output
            if tailored_cv_output.startswith("```json\n"):
                tailored_cv_output = tailored_cv_output[len("```json\n"):]
            elif tailored_cv_output.startswith("```json"): # Handle if no newline after json
                tailored_cv_output = tailored_cv_output[len("```json"):]
            elif tailored_cv_output.startswith("```"): # Handle if just ```
                tailored_cv_output = tailored_cv_output[len("```"):]

            if tailored_cv_output.endswith("\n```"):
                tailored_cv_output = tailored_cv_output[:-len("\n```")]
            elif tailored_cv_output.endswith("```"): # Handle if no newline before trailing ```
                tailored_cv_output = tailored_cv_output[:-len("```")]

            tailored_cv_output = tailored_cv_output.strip() # Clean any surrounding whitespace

            print("\n--- Step 4: Tailored CV Output ---")
            print(tailored_cv_output) # This will now print the cleaned output
            
            # Directly attempt to generate PDF from tailored CV
            print("\nAttempting to generate PDF from tailored CV...")
            try:
                generate_cv_pdf_from_json_string(tailored_cv_output)
                # Assuming generate_cv_pdf_from_json_string prints its own success/failure messages
                # print(f"Successfully generated PDF: tailored_cv.pdf") # Or get filename from the function
            except Exception as e_pdf:
                print(f"Error during PDF generation process: {e_pdf}")
        else:
            print("\nFailed to get tailored CV from API.")

    else:
        print("\nSkipping processing as either CV or Job Description is missing.")

    print("\n--- Program End ---")

if __name__ == "__main__":
    main()
